import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:  # Package import: ``from scripts import build_paper_docx``.
    from .omml_math import append_display_equation, append_inline_equation
except ImportError:  # Direct CLI execution: ``python scripts/build_paper_docx.py``.
    from omml_math import append_display_equation, append_inline_equation


DEFAULT_SECTIONS = [
    ("1. Introduction", ["State the research question, setting, empirical strategy, and main finding."]),
    ("2. Data", ["Describe sample construction, variables, transformations, and summary statistics."]),
    ("3. Empirical Strategy", ["Present the baseline equation, estimator, assumptions, and inference approach."]),
    ("4. Results", ["Interpret estimates with economic magnitude and uncertainty."]),
    ("5. Robustness", ["Group robustness checks by threat: measurement, sample, specification, inference, and identification."]),
    ("6. Conclusion", ["Summarize what the evidence shows and what remains uncertain."]),
]


def _require_keys(value, required, context):
    missing = [key for key in required if key not in value]
    if missing:
        raise ValueError(f"{context} is missing required field(s): {', '.join(missing)}")


def _reject_unknown_keys(value, allowed, context):
    unknown = sorted(set(value) - set(allowed))
    if unknown:
        raise ValueError(f"{context} has unknown field(s): {', '.join(unknown)}")


def _validate_runs(runs, context):
    if not isinstance(runs, list):
        raise ValueError(f"{context}.runs must be an array")
    allowed = {"text", "bold", "italic", "superscript", "subscript", "size"}
    for index, run in enumerate(runs):
        if isinstance(run, str):
            continue
        if not isinstance(run, dict):
            raise ValueError(f"{context}.runs[{index}] must be a string or object")
        _require_keys(run, ["text"], f"{context}.runs[{index}]")
        _reject_unknown_keys(run, allowed, f"{context}.runs[{index}]")


def _validate_table_cell(cell, context):
    if not isinstance(cell, dict):
        return
    allowed = {"text", "math", "bold", "italic", "size"}
    _reject_unknown_keys(cell, allowed, context)
    if "text" not in cell and "math" not in cell:
        raise ValueError(f"{context} requires 'text' or 'math'")
    if isinstance(cell.get("math"), bool) and cell.get("math") and "text" not in cell:
        raise ValueError(f"{context} with math=true also requires 'text'")


def _validate_table_merge(merge, context, row_count, column_count):
    def index(value, name, upper):
        if not isinstance(value, int) or isinstance(value, bool):
            raise ValueError(f"{context}.{name} must be an integer")
        if not 0 <= value < upper:
            raise ValueError(
                f"{context}.{name}={value} is outside the table range 0..{upper - 1}"
            )
        return value

    if isinstance(merge, dict):
        _reject_unknown_keys(
            merge,
            {"row", "start_col", "end_col", "col", "start_row", "end_row"},
            context,
        )
        vertical = "col" in merge
        required = ["col", "start_row", "end_row"] if vertical else [
            "row", "start_col", "end_col"
        ]
        _require_keys(merge, required, context)
        if vertical:
            index(merge["col"], "col", column_count)
            start = index(merge["start_row"], "start_row", row_count)
            end = index(merge["end_row"], "end_row", row_count)
        else:
            index(merge["row"], "row", row_count)
            start = index(merge["start_col"], "start_col", column_count)
            end = index(merge["end_col"], "end_col", column_count)
    elif isinstance(merge, list):
        if len(merge) not in (3, 4) or (len(merge) == 4 and merge[3] != "vertical"):
            raise ValueError(
                f"{context} must contain three indices and optional fourth value 'vertical'"
            )
        vertical = len(merge) == 4
        if vertical:
            index(merge[0], "col", column_count)
            start = index(merge[1], "start_row", row_count)
            end = index(merge[2], "end_row", row_count)
        else:
            index(merge[0], "row", row_count)
            start = index(merge[1], "start_col", column_count)
            end = index(merge[2], "end_col", column_count)
    else:
        raise ValueError(f"{context} must be an object or array")
    if end <= start:
        raise ValueError(f"{context} end index must be greater than start index")


def _validate_content_item(item, context):
    if isinstance(item, str):
        return
    if not isinstance(item, dict):
        raise ValueError(f"{context} must be a string or content object")
    item_type = item.get("type")
    if item_type == "paragraph":
        allowed = {"type", "text", "runs", "size"}
        _reject_unknown_keys(item, allowed, context)
        if ("text" in item) == ("runs" in item):
            raise ValueError(f"{context} requires exactly one of 'text' or 'runs'")
        if "runs" in item:
            _validate_runs(item["runs"], context)
        return
    if item_type == "equation":
        _require_keys(item, ["text"], context)
        _reject_unknown_keys(item, {"type", "text"}, context)
        if not isinstance(item["text"], str) or not item["text"].strip():
            raise ValueError(f"{context}.text must be a non-empty formula string")
        return
    if item_type == "table":
        _require_keys(item, ["rows"], context)
        allowed = {
            "type", "caption", "rows", "header_rows", "header_merges", "merges",
            "style", "notes", "source", "column_widths_inches",
            "column_alignments", "repeat_header",
        }
        _reject_unknown_keys(item, allowed, context)
        rows = item["rows"]
        if not isinstance(rows, list) or not rows:
            raise ValueError(f"{context}.rows must be a non-empty array")
        for row_index, row in enumerate(rows):
            if not isinstance(row, list) or not row:
                raise ValueError(f"{context}.rows[{row_index}] must be a non-empty array")
            for col_index, cell in enumerate(row):
                _validate_table_cell(cell, f"{context}.rows[{row_index}][{col_index}]")
        column_count = max(len(row) for row in rows)
        for merge_field in ("header_merges", "merges"):
            merge_values = item.get(merge_field, [])
            if not isinstance(merge_values, list):
                raise ValueError(f"{context}.{merge_field} must be an array")
            for merge_index, merge in enumerate(merge_values):
                _validate_table_merge(
                    merge,
                    f"{context}.{merge_field}[{merge_index}]",
                    len(rows),
                    column_count,
                )
        if "header_rows" in item:
            header_rows = item["header_rows"]
            if not isinstance(header_rows, int) or isinstance(header_rows, bool) or not 1 <= header_rows <= len(rows):
                raise ValueError(f"{context}.header_rows must be an integer between 1 and row count")
        if item.get("style", "three-line") not in {"three-line", "grid"}:
            raise ValueError(f"{context}.style must be 'three-line' or 'grid'")
        widths = item.get("column_widths_inches")
        if widths is not None:
            if not isinstance(widths, list) or len(widths) != column_count:
                raise ValueError(f"{context}.column_widths_inches must match the table column count")
            if any(not isinstance(value, (int, float)) or isinstance(value, bool) or value <= 0 for value in widths):
                raise ValueError(f"{context}.column_widths_inches values must be positive numbers")
        alignments = item.get("column_alignments")
        if alignments is not None:
            if not isinstance(alignments, list) or len(alignments) != column_count:
                raise ValueError(f"{context}.column_alignments must match the table column count")
            if any(value not in {"left", "center", "right"} for value in alignments):
                raise ValueError(f"{context}.column_alignments values must be left, center, or right")
        if "repeat_header" in item and not isinstance(item["repeat_header"], bool):
            raise ValueError(f"{context}.repeat_header must be a boolean")
        return
    if item_type == "figure":
        _require_keys(item, ["path"], context)
        allowed = {"type", "path", "width_inches", "caption", "caption_position", "notes", "source"}
        _reject_unknown_keys(item, allowed, context)
        if not isinstance(item["path"], str) or not item["path"].strip():
            raise ValueError(f"{context}.path must be a non-empty string")
        if item.get("caption_position", "below") not in {"above", "below"}:
            raise ValueError(f"{context}.caption_position must be 'above' or 'below'")
        return
    raise ValueError(
        f"{context}.type must be one of paragraph, equation, table, or figure; got {item_type!r}"
    )


def validate_paper_spec(spec):
    """Validate the dependency-free subset documented by the JSON Schema."""
    if not isinstance(spec, dict):
        raise ValueError("paper input must be a JSON object")
    allowed = {"title", "author", "affiliation", "date", "abstract", "keywords", "sections", "references"}
    _require_keys(spec, ["title"], "paper input")
    _reject_unknown_keys(spec, allowed, "paper input")
    if not isinstance(spec["title"], str) or not spec["title"].strip():
        raise ValueError("paper input.title must be a non-empty string")
    for key in ["author", "affiliation", "date", "abstract"]:
        if key in spec and not isinstance(spec[key], str):
            raise ValueError(f"paper input.{key} must be a string")
    if "keywords" in spec:
        keywords = spec["keywords"]
        if not isinstance(keywords, str) and not (
            isinstance(keywords, list) and all(isinstance(item, str) for item in keywords)
        ):
            raise ValueError("paper input.keywords must be a string or an array of strings")

    if "sections" in spec:
        if not isinstance(spec["sections"], list):
            raise ValueError("paper input.sections must be an array")
        for section_index, section in enumerate(spec["sections"]):
            context = f"paper input.sections[{section_index}]"
            if not isinstance(section, dict):
                raise ValueError(f"{context} must be an object")
            _require_keys(section, ["heading", "content"], context)
            _reject_unknown_keys(section, {"heading", "style", "content"}, context)
            if not isinstance(section["heading"], str) or not section["heading"].strip():
                raise ValueError(f"{context}.heading must be a non-empty string")
            if not isinstance(section["content"], list):
                raise ValueError(f"{context}.content must be an array")
            for item_index, item in enumerate(section["content"]):
                _validate_content_item(item, f"{context}.content[{item_index}]")

    if "references" in spec:
        if not isinstance(spec["references"], list):
            raise ValueError("paper input.references must be an array")
        for index, reference in enumerate(spec["references"]):
            if isinstance(reference, str):
                continue
            if not isinstance(reference, dict):
                raise ValueError(f"paper input.references[{index}] must be a string or object")
            _reject_unknown_keys(reference, {"text", "runs"}, f"paper input.references[{index}]")
            if ("text" in reference) == ("runs" in reference):
                raise ValueError(f"paper input.references[{index}] requires exactly one of text or runs")
            if "runs" in reference:
                _validate_runs(reference["runs"], f"paper input.references[{index}]")
    return spec


def set_east_asia_font(r_pr, font_name):
    if not font_name:
        return
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def set_run_style(
    run,
    bold=False,
    italic=False,
    superscript=False,
    subscript=False,
    size=12,
    font_name=None,
    east_asia_font=None,
):
    run.bold = bool(bold)
    run.italic = bool(italic)
    run.font.superscript = bool(superscript)
    run.font.subscript = bool(subscript)
    # Leave fonts unset by default so supplied template styles remain
    # authoritative. Explicit overrides set Latin/East Asian fonts.
    if font_name:
        run.font.name = font_name
    if east_asia_font:
        set_east_asia_font(run._r.get_or_add_rPr(), east_asia_font)
    run.font.size = Pt(size)


def add_styled_run(
    paragraph,
    text,
    bold=False,
    italic=False,
    superscript=False,
    subscript=False,
    size=12,
    font_name=None,
    east_asia_font=None,
):
    run = paragraph.add_run(str(text))
    set_run_style(
        run,
        bold=bold,
        italic=italic,
        superscript=superscript,
        subscript=subscript,
        size=size,
        font_name=font_name,
        east_asia_font=east_asia_font,
    )
    return run


def add_runs(paragraph, runs, default_size=12):
    for item in runs:
        if isinstance(item, str):
            add_styled_run(paragraph, item, size=default_size)
            continue
        add_styled_run(
            paragraph,
            item.get("text", ""),
            bold=item.get("bold", False),
            italic=item.get("italic", False),
            superscript=item.get("superscript", False),
            subscript=item.get("subscript", False),
            size=item.get("size", default_size),
        )


def set_cell_text(cell, value, bold=False, alignment="left"):
    """Set a table cell and render OMML only when explicitly requested.

    Plain values always remain text. Use ``{"text": "beta_1", "math": true}``
    or the shorthand ``{"math": "beta_1"}`` for Word-native math.
    """
    options = value if isinstance(value, dict) else {}
    math_option = options.get("math", False)
    if isinstance(math_option, str) and "text" not in options:
        text = math_option
        render_math = True
    else:
        text = options.get("text", value)
        render_math = bool(math_option)
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[alignment]
    if render_math:
        append_inline_equation(p, str(text))
    else:
        add_styled_run(
            p,
            text,
            bold=options.get("bold", bold),
            italic=options.get("italic", False),
            size=options.get("size", 10),
        )


def set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key, value in kwargs[edge].items():
            element.set(qn("w:{}".format(key)), str(value))


def force_style_color(style, color="000000"):
    r_pr = style.element.get_or_add_rPr()
    color_el = r_pr.find(qn("w:color"))
    if color_el is None:
        color_el = OxmlElement("w:color")
        r_pr.append(color_el)
    color_el.set(qn("w:val"), color)


def set_style_fonts(style, latin_font="Times New Roman", east_asia_font="Times New Roman"):
    """Set all OOXML font slots for an explicitly applied default style."""
    style.font.name = latin_font
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attribute in ["ascii", "hAnsi", "cs"]:
        r_fonts.set(qn(f"w:{attribute}"), latin_font)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)


def remove_style_paragraph_borders(style):
    p_pr = style.element.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is not None:
        p_pr.remove(borders)


def apply_three_line_table(table, header_rows=1):
    line = {"val": "single", "sz": "8", "space": "0", "color": "000000"}
    no_line = {"val": "nil"}
    row_count = len(table.rows)
    if row_count == 0:
        return
    header_rows = max(1, min(int(header_rows), row_count))
    header_bottom_row = header_rows - 1

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top=no_line,
                left=no_line,
                bottom=no_line,
                right=no_line,
                insideH=no_line,
                insideV=no_line,
            )

    for cell in table.rows[0].cells:
        set_cell_border(cell, top=line, left=no_line, right=no_line)
    for cell in table.rows[header_bottom_row].cells:
        set_cell_border(cell, bottom=line, left=no_line, right=no_line)
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom=line, left=no_line, right=no_line)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    set_style_fonts(normal)
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    style_specs = {
        "Title": (16, True),
        "Subtitle": (12, False),
        "Caption": (10, False),
        "Heading 1": (13, True),
        "Heading 2": (12, True),
        "Heading 3": (12, True),
    }
    for name, (size, bold) in style_specs.items():
        if name not in styles:
            continue
        style = styles[name]
        set_style_fonts(style)
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor(0, 0, 0)
        force_style_color(style)
        if name == "Title":
            remove_style_paragraph_borders(style)

    for style in styles:
        if style.name.startswith("Heading"):
            style.font.color.rgb = RGBColor(0, 0, 0)
            force_style_color(style)


def add_title_block(doc, spec):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_styled_run(title, spec.get("title", "Econometric Research Paper"), bold=True, size=16)

    for key in ["author", "affiliation", "date"]:
        if spec.get(key):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_styled_run(p, spec[key], size=12)

    if spec.get("abstract"):
        doc.add_paragraph("Abstract", style="Heading 1")
        doc.add_paragraph(spec["abstract"])

    if spec.get("keywords"):
        p = doc.add_paragraph()
        add_styled_run(p, "Keywords: ", bold=True)
        add_styled_run(p, ", ".join(spec["keywords"]) if isinstance(spec["keywords"], list) else str(spec["keywords"]))


def add_small_note(doc, label, value):
    if not value:
        return
    text = " ".join(str(item) for item in value) if isinstance(value, list) else str(value)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    add_styled_run(p, f"{label}: ", bold=True, size=9)
    add_styled_run(p, text, size=9)


def add_table(doc, table_spec):
    rows = table_spec.get("rows", [])
    if not rows:
        return
    caption = table_spec.get("caption")
    if caption:
        p = doc.add_paragraph()
        add_styled_run(p, caption, bold=True, size=10)
    col_count = max(len(row) for row in rows)
    column_alignments = table_spec.get(
        "column_alignments", ["left"] + ["center"] * max(0, col_count - 1)
    )
    table = doc.add_table(rows=len(rows), cols=col_count)
    for i, row in enumerate(rows):
        for j in range(col_count):
            value = row[j] if j < len(row) else ""
            set_cell_text(
                table.cell(i, j), value, bold=(i == 0), alignment=column_alignments[j]
            )
    apply_table_merges(table, table_spec.get("merges", []) + table_spec.get("header_merges", []))
    header_rows = table_spec.get("header_rows")
    if header_rows is None:
        header_merges = table_spec.get("header_merges", [])
        horizontal_header_rows = [
            int(merge.get("row", 0))
            for merge in header_merges
            if isinstance(merge, dict) and "col" not in merge
        ]
        # A spanning group-header row normally has a label row below it.
        # Callers can set ``header_rows`` to override this inference.
        header_rows = min(max(horizontal_header_rows, default=-1) + 2, len(rows))
        header_rows = max(header_rows, 1)
    if table_spec.get("style", "three-line") == "grid":
        table.style = "Table Grid"
    else:
        apply_three_line_table(table, header_rows=header_rows)
    if table_spec.get("repeat_header", True):
        for row in table.rows[:header_rows]:
            set_repeat_table_header(row)
    widths = table_spec.get("column_widths_inches")
    if widths:
        table.autofit = False
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(float(width))
    add_small_note(doc, "Notes", table_spec.get("notes"))
    add_small_note(doc, "Source", table_spec.get("source"))


def clean_merged_cell_paragraphs(cell):
    """Remove empty paragraphs introduced by python-docx cell merges."""
    for paragraph in list(cell.paragraphs):
        if len(cell.paragraphs) <= 1:
            break
        has_content = bool(paragraph.text.strip()) or bool(
            paragraph._p.xpath(".//m:oMath | .//w:drawing | .//w:object")
        )
        if not has_content:
            paragraph._element.getparent().remove(paragraph._element)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def apply_table_merges(table, merges):
    for merge in merges:
        if isinstance(merge, dict):
            if "col" in merge:
                col = int(merge["col"])
                start_row = int(merge.get("start_row", merge.get("start", 0)))
                end_row = int(merge.get("end_row", merge.get("end", start_row)))
                max_row = len(table.rows) - 1
                start_row = max(0, min(start_row, max_row))
                end_row = max(0, min(end_row, max_row))
                if end_row > start_row:
                    merged = table.cell(start_row, col).merge(table.cell(end_row, col))
                    clean_merged_cell_paragraphs(merged)
            else:
                row = int(merge.get("row", 0))
                start_col = int(merge.get("start_col", merge.get("start", 0)))
                end_col = int(merge.get("end_col", merge.get("end", start_col)))
                max_col = len(table.columns) - 1
                start_col = max(0, min(start_col, max_col))
                end_col = max(0, min(end_col, max_col))
                if end_col > start_col:
                    merged = table.cell(row, start_col).merge(table.cell(row, end_col))
                    clean_merged_cell_paragraphs(merged)
        else:
            if len(merge) >= 4 and str(merge[3]).lower() == "vertical":
                col, start_row, end_row = [int(v) for v in merge[:3]]
                max_row = len(table.rows) - 1
                start_row = max(0, min(start_row, max_row))
                end_row = max(0, min(end_row, max_row))
                if end_row > start_row:
                    merged = table.cell(start_row, col).merge(table.cell(end_row, col))
                    clean_merged_cell_paragraphs(merged)
            else:
                row, start_col, end_col = [int(value) for value in merge[:3]]
                max_col = len(table.columns) - 1
                start_col = max(0, min(start_col, max_col))
                end_col = max(0, min(end_col, max_col))
                if end_col > start_col:
                    merged = table.cell(row, start_col).merge(table.cell(row, end_col))
                    clean_merged_cell_paragraphs(merged)


def add_figure(doc, figure_spec):
    path = Path(figure_spec.get("path", ""))
    if not path.exists():
        raise FileNotFoundError(f"Figure file not found: {path}")

    width = Inches(float(figure_spec.get("width_inches", 5.8)))
    caption = figure_spec.get("caption")
    if caption and figure_spec.get("caption_position") == "above":
        p = doc.add_paragraph()
        add_styled_run(p, caption, bold=True, size=10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=width)

    if caption and figure_spec.get("caption_position", "below") != "above":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_styled_run(p, caption, bold=True, size=10)

    add_small_note(doc, "Notes", figure_spec.get("notes"))
    add_small_note(doc, "Source", figure_spec.get("source"))


def add_paragraph_item(doc, item):
    p = doc.add_paragraph()
    if item.get("runs"):
        add_runs(p, item["runs"], default_size=item.get("size", 12))
    else:
        add_styled_run(p, item.get("text", ""), size=item.get("size", 12))
    return p


def add_section(doc, section):
    doc.add_paragraph(section.get("heading", "Section"), style=section.get("style", "Heading 1"))
    for item in section.get("content", []):
        if isinstance(item, str):
            doc.add_paragraph(item)
        elif item.get("type") == "paragraph":
            add_paragraph_item(doc, item)
        elif item.get("type") == "equation":
            append_display_equation(doc, item.get("text", ""))
        elif item.get("type") == "table":
            add_table(doc, item)
        elif item.get("type") == "figure":
            add_figure(doc, item)


def build_docx(spec, out_path, template=None, apply_default_style=None):
    validate_paper_spec(spec)
    doc = Document(template) if template else Document()
    if apply_default_style is None:
        apply_default_style = not bool(template)
    if apply_default_style:
        style_document(doc)
    add_title_block(doc, spec)

    sections = spec.get("sections")
    if not sections:
        sections = [{"heading": h, "content": [{"type": "paragraph", "text": p} for p in ps]} for h, ps in DEFAULT_SECTIONS]
    for section in sections:
        add_section(doc, section)

    if spec.get("references"):
        doc.add_paragraph("References", style="Heading 1")
        for ref in spec["references"]:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            if isinstance(ref, dict) and ref.get("runs"):
                add_runs(p, ref["runs"], default_size=12)
            else:
                add_styled_run(p, ref.get("text", "") if isinstance(ref, dict) else ref)

    doc.save(out_path)


def build_style_template(out_path):
    doc = Document()
    style_document(doc)
    doc.save(out_path)


def main():
    parser = argparse.ArgumentParser(description="Build a standard economics/management paper DOCX from JSON.")
    parser.add_argument("input_json")
    parser.add_argument("output_docx")
    parser.add_argument("--template")
    parser.add_argument("--apply-default-style", action="store_true", help="Apply the built-in Times New Roman layout even when a template is supplied.")
    parser.add_argument("--make-template", action="store_true", help="Ignore input content and create a reusable style template.")
    args = parser.parse_args()

    if args.make_template:
        build_style_template(args.output_docx)
        return
    else:
        spec = json.loads(Path(args.input_json).read_text(encoding="utf-8"))

    apply_default_style = True if args.apply_default_style else None
    try:
        build_docx(spec, args.output_docx, template=args.template, apply_default_style=apply_default_style)
    except (ValueError, FileNotFoundError) as exc:
        raise SystemExit(f"Error: {exc}")


if __name__ == "__main__":
    main()
