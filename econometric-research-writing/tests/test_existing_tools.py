import json
import os
from pathlib import Path
import subprocess
import sys
import tempfile
import textwrap
import unittest
from unittest import mock
import zipfile

import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from lxml import etree


ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

# This package-style import is itself a regression test for the relative
# omml_math import fallback in build_paper_docx.
from scripts import build_paper_docx as builder
from scripts import check_docx_integrity as integrity
from scripts import omml_math
from scripts import profile_econ_dataset as profiler
from scripts import verify_references as verifier


class ProfileDatasetTests(unittest.TestCase):
    def test_xlsx_without_sheet_reads_first_sheet_dataframe(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "sample.xlsx"
            pd.DataFrame({"firm_id": [1, 2], "year": [2020, 2021]}).to_excel(path, index=False)
            frame = profiler.read_data(path)
            self.assertIsInstance(frame, pd.DataFrame)
            self.assertEqual(frame.shape, (2, 2))

    def test_tiny_sample_markdown_formats_missing_standard_deviation(self):
        frame = pd.DataFrame({"value": [1.0]})
        report = profiler.build_report(Path("tiny.csv"), frame, {})
        markdown = profiler.markdown_report(report)
        self.assertIn("sd n/a", markdown)

    def test_balanced_panel_requires_complete_unit_time_grid(self):
        equal_counts_but_not_grid = pd.DataFrame({
            "unit": ["a", "a", "b", "b"],
            "time": [1, 2, 2, 3],
        })
        check = profiler.panel_checks(equal_counts_but_not_grid, {"unit": "unit", "time": "time"})
        self.assertFalse(check["balanced_candidate"])
        self.assertEqual(check["observed_grid_cells"], 4)
        self.assertEqual(check["expected_grid_cells"], 6)

        full_grid = pd.DataFrame({
            "unit": ["a", "a", "b", "b"],
            "time": [1, 2, 1, 2],
        })
        check = profiler.panel_checks(full_grid, {"unit": "unit", "time": "time"})
        self.assertTrue(check["complete_unit_time_grid"])

    def test_name_hints_remain_provisional_and_do_not_draw_causal_map(self):
        frame = pd.DataFrame({
            "firm_id": [1, 2],
            "year": [2020, 2020],
            "treatment": [0, 1],
            "outcome": [2.0, 3.0],
        })
        report = profiler.build_report(Path("roles.csv"), frame, {})
        self.assertEqual(set(report["analysis_roles"]), {"unit", "time"})
        self.assertEqual(report["role_sources"]["unit"], "provisional_name_hint")
        markdown = profiler.markdown_report(report)
        self.assertNotIn("Causal Pathway", markdown)
        self.assertNotIn("-->", markdown)

    def test_outcome_treatment_correlation_is_not_called_multicollinearity(self):
        frame = pd.DataFrame({
            "outcome": [2.0, 4.0, 6.0, 8.0],
            "treatment": [1.0, 2.0, 3.0, 4.0],
            "control": [4.0, 1.0, 3.0, 2.0],
        })
        report = profiler.build_report(
            Path("roles.csv"),
            frame,
            {"outcome": "outcome", "treatment": "treatment"},
        )
        self.assertFalse(any("multicollinearity" in item for item in report["warnings"]))

    def test_declared_correlated_regressors_trigger_targeted_warning(self):
        frame = pd.DataFrame({
            "outcome": [0.0, 1.0, 0.0, 1.0],
            "treatment": [1.0, 2.0, 3.0, 4.0],
            "control": [2.0, 4.0, 6.0, 8.0],
        })
        report = profiler.build_report(
            Path("roles.csv"),
            frame,
            {
                "outcome": "outcome",
                "treatment": "treatment",
                "controls": ["control"],
            },
        )
        self.assertTrue(
            any("declared regressors" in item for item in report["warnings"])
        )


class ReferenceVerificationTests(unittest.TestCase):
    @staticmethod
    def crossref_item(title="政策效果研究", family="Doe", year=2024):
        return {
            "title": [title],
            "author": [{"given": "Jane", "family": family}],
            "issued": {"date-parts": [[year]]},
            "container-title": ["Journal of Applied Economics"],
            "DOI": "10.1234/example",
            "URL": "https://doi.org/10.1234/example",
            "type": "journal-article",
        }

    def test_unicode_title_is_preserved_for_matching(self):
        self.assertEqual(verifier.normalize_title("政策—效果：研究"), "政策 效果 研究")
        self.assertEqual(verifier.title_similarity("政策效果研究", "政策效果研究"), 1.0)

    def test_title_only_result_is_candidate_not_verified(self):
        with mock.patch.object(verifier, "crossref_by_title", return_value=[self.crossref_item()]):
            result = verifier.verify_reference({"title": "政策效果研究"}, sleep_seconds=0)
        self.assertEqual(result["status"], "candidate_match")

    def test_author_and_year_corroborate_metadata_match(self):
        ref = {"title": "政策效果研究", "authors": ["Jane Doe"], "year": 2024}
        with mock.patch.object(verifier, "crossref_by_title", return_value=[self.crossref_item()]):
            result = verifier.verify_reference(ref, sleep_seconds=0)
        self.assertEqual(result["status"], "verified_by_metadata")
        self.assertEqual(result["field_scores"]["authors"], 1.0)
        self.assertEqual(result["field_scores"]["year"], 1.0)

    def test_conflicting_author_is_metadata_mismatch(self):
        ref = {"title": "政策效果研究", "authors": ["Jane Smith"], "year": 2024}
        with mock.patch.object(verifier, "crossref_by_title", return_value=[self.crossref_item()]):
            result = verifier.verify_reference(ref, sleep_seconds=0)
        self.assertEqual(result["status"], "metadata_mismatch")
        self.assertIn("authors", result["warning"])


class OmmlTests(unittest.TestCase):
    def test_text_command_accepts_trailing_subscript(self):
        element = omml_math.omath_text(r"\text{ATT}_{g,t}")
        xml = etree.tostring(element, encoding="unicode")
        self.assertIn("sSub", xml)
        self.assertIn("ATT", xml)

    def test_unsupported_or_malformed_syntax_fails(self):
        with self.assertRaises(omml_math.UnsupportedFormulaSyntax):
            omml_math.omath_text(r"\unknown{x}")
        with self.assertRaises(omml_math.UnsupportedFormulaSyntax):
            omml_math.omath_text("x_")
        with self.assertRaises(omml_math.UnsupportedFormulaSyntax):
            omml_math.omath_text(r"\frac{x}{y")


class DocxBuilderAndIntegrityTests(unittest.TestCase):
    def test_example_contract_validates_and_builds(self):
        schema = json.loads((ROOT / "schemas/paper-docx.schema.json").read_text(encoding="utf-8"))
        spec = json.loads((ROOT / "examples/paper-docx.example.json").read_text(encoding="utf-8"))
        self.assertEqual(schema["$schema"], "https://json-schema.org/draft/2020-12/schema")
        self.assertIn("title", schema["required"])
        builder.validate_paper_spec(spec)
        with tempfile.TemporaryDirectory() as tmp:
            output = Path(tmp) / "example.docx"
            builder.build_docx(spec, output)
            report = integrity.inspect_docx(output)
            built = Document(output)
            table = built.tables[0]
            repeats = len(table.rows[0]._tr.xpath("./w:trPr/w:tblHeader"))
            numeric_alignment = table.cell(1, 1).paragraphs[0].alignment
            autofit = table.autofit
        self.assertGreaterEqual(report["omml_count"], 2)
        self.assertEqual(report["table_count"], 1)
        self.assertEqual(report["three_line_table_count"], 1)
        self.assertEqual(repeats, 1)
        self.assertEqual(numeric_alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertFalse(autofit)

    def test_invalid_contract_is_rejected(self):
        with self.assertRaisesRegex(ValueError, "title"):
            builder.validate_paper_spec({"sections": []})
        with self.assertRaisesRegex(ValueError, "type must be one of"):
            builder.validate_paper_spec({
                "title": "Bad",
                "sections": [{"heading": "X", "content": [{"type": "unknown"}]}],
            })
        with self.assertRaisesRegex(ValueError, "missing required field.*rows"):
            builder.validate_paper_spec({
                "title": "Bad table",
                "sections": [{"heading": "X", "content": [{"type": "table"}]}],
            })
        with self.assertRaisesRegex(ValueError, "outside the table range"):
            builder.validate_paper_spec({
                "title": "Bad merge",
                "sections": [{
                    "heading": "X",
                    "content": [{
                        "type": "table",
                        "rows": [["A", "B"]],
                        "header_merges": [
                            {"row": 99, "start_col": 0, "end_col": 1}
                        ],
                    }],
                }],
            })

    def test_table_math_is_explicit_and_multirow_rule_is_detected(self):
        spec = {
            "title": "Table test",
            "sections": [{
                "heading": "Results",
                "content": [{
                    "type": "table",
                    "header_rows": 2,
                    "rows": [
                        ["", "Outcome", ""],
                        ["Variable", "(1)", "(2)"],
                        ["firm_id", {"text": "beta_1", "math": True}, "0.1"],
                    ],
                    "header_merges": [{"row": 0, "start_col": 1, "end_col": 2}],
                }],
            }],
        }
        with tempfile.TemporaryDirectory() as tmp:
            output = Path(tmp) / "table.docx"
            builder.build_docx(spec, output)
            report = integrity.inspect_docx(output)
            built = Document(output)
            with zipfile.ZipFile(output) as archive:
                document_xml = archive.read("word/document.xml").decode("utf-8")
        self.assertIn("firm_id", document_xml)
        self.assertEqual(report["omml_count"], 1)
        self.assertEqual(report["three_line_table_count"], 1)
        self.assertEqual(report["table_border_summaries"][0]["header_separator_rows"], [1])
        self.assertFalse(report["table_border_summaries"][0]["has_vertical_borders"])
        self.assertEqual(len(built.tables[0].cell(0, 1).paragraphs), 1)
        self.assertEqual(
            integrity.assertion_failures(
                report, min_omml=1, min_tables=1, require_three_line_tables=True, forbid_grid_tables=True
            ),
            [],
        )
        self.assertTrue(integrity.assertion_failures(report, min_media=1))

    def test_grid_table_can_repeat_header(self):
        spec = {
            "title": "Grid test",
            "sections": [{
                "heading": "Appendix",
                "content": [{
                    "type": "table",
                    "style": "grid",
                    "header_rows": 1,
                    "repeat_header": True,
                    "rows": [["Variable", "Value"], ["x", "1"]],
                }],
            }],
        }
        with tempfile.TemporaryDirectory() as tmp:
            output = Path(tmp) / "grid.docx"
            builder.build_docx(spec, output)
            built = Document(output)
            report = integrity.inspect_docx(output)
        self.assertEqual(report["grid_table_count"], 1)
        self.assertEqual(len(built.tables[0].rows[0]._tr.xpath("./w:trPr/w:tblHeader")), 1)

    def test_supplied_template_fonts_are_inherited_not_overridden(self):
        with tempfile.TemporaryDirectory() as tmp:
            template_path = Path(tmp) / "template.docx"
            output_path = Path(tmp) / "output.docx"
            template = Document()
            normal = template.styles["Normal"]
            normal.font.name = "Arial"
            r_pr = normal.element.get_or_add_rPr()
            r_fonts = r_pr.find(qn("w:rFonts"))
            if r_fonts is None:
                r_fonts = OxmlElement("w:rFonts")
                r_pr.insert(0, r_fonts)
            r_fonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
            template.save(template_path)

            builder.build_docx({"title": "Template test", "sections": []}, output_path, template=template_path)
            output = Document(output_path)
            output_fonts = output.styles["Normal"].element.get_or_add_rPr().find(qn("w:rFonts"))
            first_run_fonts = output.paragraphs[-7].runs[0]._r.get_or_add_rPr().find(qn("w:rFonts"))
        self.assertEqual(output.styles["Normal"].font.name, "Arial")
        self.assertEqual(output_fonts.get(qn("w:eastAsia")), "Hiragino Sans GB")
        self.assertIsNone(first_run_fonts)

    def test_default_template_core_styles_are_clean(self):
        report = integrity.inspect_docx(ROOT / "assets/econ-paper-template.docx")
        self.assertEqual(
            integrity.assertion_failures(report, require_clean_core_styles=True),
            [],
        )

    def test_nonblack_direct_heading_formatting_is_detected(self):
        with tempfile.TemporaryDirectory() as tmp:
            output = Path(tmp) / "blue-heading.docx"
            document = Document()
            paragraph = document.add_paragraph(style="Heading 1")
            run = paragraph.add_run("Direct blue heading")
            run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
            document.save(output)
            report = integrity.inspect_docx(output)
        self.assertTrue(report["nonblack_heading_direct_formatting"])
        self.assertTrue(
            any(
                "direct formatting" in failure
                for failure in integrity.assertion_failures(
                    report, fail_on_nonblack_headings=True
                )
            )
        )


class RenderWrapperTests(unittest.TestCase):
    def test_missing_input_is_nonzero(self):
        script = ROOT / "scripts/render_validate_docx.py"
        result = subprocess.run(
            [sys.executable, str(script), "/definitely/missing.docx", "--output-dir", "/tmp/render-missing"],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )
        self.assertNotEqual(result.returncode, 0)
        self.assertIn("does not exist", result.stderr + result.stdout)

    def test_stale_pages_are_removed_before_render(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp = Path(tmp)
            input_path = tmp / "input.docx"
            input_path.write_bytes(b"fake input accepted by fake renderer")
            output_dir = tmp / "render"
            output_dir.mkdir()
            stale = output_dir / "page-99.png"
            stale.write_bytes(b"stale")
            renderer = tmp / "fake_renderer.py"
            renderer.write_text(textwrap.dedent("""
                import argparse
                from pathlib import Path
                parser = argparse.ArgumentParser()
                parser.add_argument('docx')
                parser.add_argument('--output_dir', required=True)
                parser.add_argument('--emit_pdf', action='store_true')
                args = parser.parse_args()
                out = Path(args.output_dir)
                out.mkdir(parents=True, exist_ok=True)
                (out / 'page-1.png').write_bytes(b'new')
            """), encoding="utf-8")
            env = os.environ.copy()
            env["DOCX_RENDER_SCRIPT"] = str(renderer)
            env["PYTHON"] = sys.executable
            result = subprocess.run(
                [
                    sys.executable,
                    str(ROOT / "scripts/render_validate_docx.py"),
                    str(input_path),
                    "--output-dir",
                    str(output_dir),
                ],
                env=env,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
            )
            self.assertEqual(result.returncode, 0, result.stderr)
            self.assertFalse(stale.exists())
            self.assertEqual((output_dir / "page-1.png").read_bytes(), b"new")


if __name__ == "__main__":
    unittest.main()
